from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '游明朝'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
style.font.size = Pt(10)

# ────────────────────────────────────────
# ユーティリティ
# ────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6E)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 2:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, indent=0, bold=False, red=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.bold = bold
    if red:
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return p

def simple_table(headers, rows, col_widths, header_bg='1A376E', alt_bg='EBF3FB'):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        c = hdr_row.cells[i]
        c.width = Cm(w)
        set_cell_bg(c, header_bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        bg = alt_bg if ri % 2 == 1 else 'FFFFFF'
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            c = row.cells[ci]
            c.width = Cm(w)
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if '⚠' in str(val):
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return tbl

# ════════════════════════════════════════
# 冒頭（既存規程スタイルに統一）
# ════════════════════════════════════════
doc.add_paragraph()

# 会社名
p0 = doc.add_paragraph()
p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
r0 = p0.add_run('株式会社サキュレ　社内規程')
r0.bold = True
r0.font.size = Pt(11)
r0.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# 規程名
p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p1.add_run('タイミー（スポットワーク）活用に関する社内規程')
r1.bold = True
r1.font.size = Pt(20)
r1.font.color.rgb = RGBColor(0x1A, 0x37, 0x6E)
p1.paragraph_format.space_after = Pt(12)

# 文書管理テーブル（既存規程と同スタイル）
meta_tbl = doc.add_table(rows=1, cols=4)
meta_tbl.style = 'Table Grid'
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('文書番号\nMGT-002', '制定日\n令和8年6月3日', '改定日\n―', '管理部門\n総務・人事部門'),
]
meta_widths = [3.5, 4.0, 3.5, 4.5]
for ci, (val, w) in enumerate(zip(meta_data[0], meta_widths)):
    c = meta_tbl.rows[0].cells[ci]
    c.width = Cm(w)
    set_cell_bg(c, 'D9E1F2')
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(val)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1A, 0x37, 0x6E)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_paragraph()

# ════════════════════════════════════════
# 第1条〜第2条
# ════════════════════════════════════════
heading('第１条（目的）', level=1)
body('本規程は、株式会社サキュレ（以下「当社」という）が、スポットワーカーマッチングサービス「タイミー」（以下「タイミー」という）を活用して短期・スポット人員を確保する際の申請手続き、安全管理、業務範囲、費用処理および個人情報の取り扱いについて定めることを目的とする。')

heading('第２条（適用範囲）', level=1)
body('本規程は、タイミーを通じてスポットワーカーを受け入れるすべての部門・拠点に適用する。')

# ════════════════════════════════════════
# 第3条　利用可能業務・禁止業務
# ════════════════════════════════════════
heading('第３条（利用可能業務および禁止業務）', level=1)

heading('（１）利用可能業務', level=2)
body('タイミーを通じたスポットワーカーに従事させることができる業務は以下のとおりとする。', indent=0.5)
ok_rows = [
    ['①', '一般物品の仕分け・梱包・積み下ろし（廃棄物を含まないもの）'],
    ['②', 'リユース品の検品・清掃・梱包補助'],
    ['③', '事務所内の清掃・軽作業'],
    ['④', '倉庫内ピッキング補助'],
    ['⑤', 'イベント・繁忙期における一般業務補助'],
]
simple_table(['No.', '業務内容'], ok_rows, [1.0, 14.5])

heading('（２）禁止業務　⚠️ 厳守', level=2)
body('下記業務へのスポットワーカーの従事は、法令上・安全上の理由により禁止する。', indent=0.5, bold=True, red=True)
ng_rows = [
    ['①', '産業廃棄物の収集・運搬・積み込み・仕分け業務全般', '廃棄物処理法上、許可業者従業員のみ従事可'],
    ['②', '一般廃棄物の処理・運搬補助', '同上'],
    ['③', '車両の運転（社有車・リース車含む）', '事故リスク・保険適用外'],
    ['④', '金銭・貴重品の取り扱い', '不正リスク'],
    ['⑤', '顧客との直接交渉・商談', '信頼関係保護'],
]
simple_table(['No.', '禁止業務', '理由'], ng_rows, [0.8, 8.5, 6.2], header_bg='C00000')

# ════════════════════════════════════════
# 第4条　申請手続き
# ════════════════════════════════════════
heading('第４条（利用申請手続き）', level=1)

heading('（１）申請フロー', level=2)
flow_rows = [
    ['STEP 1', '利用依頼',  '現場責任者',         '業務内容・必要人数・日時・現場住所をTeamsで配車係へ提出（原則作業日の1週間前まで）'],
    ['STEP 2', '承認',      '配車係',              '業務内容・必要人数・日時・場所を確認し承認または差し戻し'],
    ['STEP 3', '求人掲載',  '配車係',              'タイミー管理者アプリ上で求人を作成・公開'],
    ['STEP 4', '当日受入',  '現場責任者or現場担当者', '本人確認・安全教育の実施・業務説明'],
    ['STEP 5', '終了処理',  '配車係',              '業務完了報告を受けてタイミー管理者アプリにて評価入力・実績記録'],
    ['STEP 6', '費用申請',  '経理',                '費用をタイミー請求書に基づき経費処理（勘定：外注費）'],
]
simple_table(['', 'フロー', '担当', '内容'], flow_rows, [1.0, 2.0, 3.5, 9.0])

heading('（２）申請時の必須記載事項', level=2)
body('申請時には以下の項目を必ず記載すること。', indent=0.5)
req_rows = [
    ['①', '利用日時',          '開始・終了時刻を明記（例：2026年○月○日　9:00〜12:00）'],
    ['②', '勤務場所（現地住所）', '番地まで記載。複数拠点の場合は全拠点を列記'],
    ['③', '必要人数',          '最小・最大人数を明記（タイミー上での募集枠数と一致させること）'],
    ['④', '業務内容',          '第３条（１）の利用可能業務から選択し、具体的に記載'],
    ['⑤', '服装・持ち物',      '安全靴・手袋の要否など'],
]
simple_table(['No.', '項目', '記載内容'], req_rows, [0.8, 3.5, 11.2])

# ════════════════════════════════════════
# 第5条　当日受け入れ手順
# ════════════════════════════════════════
heading('第５条（当日の受け入れ手順）', level=1)
body('現場責任者・現場担当者は、スポットワーカーの業務開始前に以下をすべて実施しなければならない。')
day_rows = [
    ['①', '本人確認',          '配車係から氏名を連絡受けて本人と照合。身分証の提示を求めてはならない（タイミー規約上の禁止事項）'],
    ['②', '⚠️ 安全教育（必須）', '入構ルール・危険箇所・緊急時避難経路・体調不良時の報告先を口頭説明'],
    ['③', '業務説明',          '作業手順・禁止行為（廃棄物エリア立ち入り禁止等）を明示'],
    ['④', '保護具の貸与',      '必要に応じて安全靴・手袋・ヘルメット等を貸与し記録'],
    ['⑤', '業務範囲の明示',    '廃棄物保管エリア・車両エリアへの立ち入り禁止区域を物理的に案内'],
]
simple_table(['No.', '項目', '内容'], day_rows, [0.8, 3.2, 11.5])

# ════════════════════════════════════════
# 第6条　安全管理
# ════════════════════════════════════════
heading('第６条（安全管理・労災対応）', level=1)
body('１　スポットワーカーは当社の管理下で業務に従事するため、業務中の労働災害については、タイミー社が加入する傷害保険が適用される。ただし、当社の安全配慮義務は免除されない。')
body('２　業務中に事故・ケガが発生した場合は、直ちに以下の手順をとること。')
body('　　（ア）応急処置・救急要請（119番）', indent=1)
body('　　（イ）配車係および総務部門へ速報', indent=1)
body('　　（ウ）タイミーサポートへ連絡（タイミーアプリ内）', indent=1)
body('　　（エ）事故報告書（別紙１）の作成・提出（発生当日中）', indent=1)
body('３　無断立ち入りや業務命令に従わない場合、即時業務終了とし、その旨をタイミーへ報告する。', bold=True, red=True)

# ════════════════════════════════════════
# 第7条　個人情報
# ════════════════════════════════════════
heading('第７条（個人情報・機密情報の取り扱い）', level=1)
body('１　スポットワーカーの氏名・顔写真等の個人情報は、タイミーのシステム上でのみ管理し、当社内で複写・保存・共有することを禁止する。')
body('２　スポットワーカーを顧客情報、社内財務情報に接触させてはならない。')

# ════════════════════════════════════════
# 第8条　費用処理
# ════════════════════════════════════════
heading('第８条（費用処理）', level=1)
body('１　タイミー利用料（ワーカー報酬＋サービス手数料）は「外注費」として計上する。')
body('２　月次締めでタイミー社より発行される請求書を総務部門が確認し、経理部門へ申請する。')
body('３　予算外の緊急利用については、社長の事前承認および事後の稟議報告を要する。')
cost_rows = [
    ['ワーカー報酬',   'タイミーが設定する時給 × 実働時間', '外注費'],
    ['サービス手数料', 'ワーカー報酬の30%（税別）',         '外注費'],
    ['傷害保険料',     'タイミー負担（当社負担なし）',      '－'],
]
simple_table(['費用区分', '計算方法', '勘定科目'], cost_rows, [3.5, 7.5, 4.5])

# ════════════════════════════════════════
# 第9条　違反時の対応
# ════════════════════════════════════════
heading('第９条（規程違反時の対応）', level=1)
body('本規程に違反した場合、以下の措置を講じることがある。')
body('　・禁止業務に従事させた場合：当該拠点のタイミー利用資格の停止', indent=0.5, bold=True, red=True)
body('　・申請手続き未履行の場合：指導・再発防止策の提出', indent=0.5)
body('　・個人情報の無断保存等：情報管理規程に基づく懲戒手続き', indent=0.5)

# ════════════════════════════════════════
# 附則
# ════════════════════════════════════════
heading('附　則', level=1)
body('本規程は、令和8年6月3日より施行する。')
body('改定が必要な場合は、総務・人事部門が原案を作成し、代表取締役の承認を得て改定する。')

# ════════════════════════════════════════
# 別紙１：事故報告書
# ════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('【別紙１】　スポットワーカー事故報告書')
r.bold = True; r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1A, 0x37, 0x6E)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.add_run('報告日：　　　年　　月　　日').font.size = Pt(10)

form2_rows = [
    ('発生日時',          '　　　年　　月　　日　　時　　分'),
    ('発生場所（住所）',  '〒　　　－　　　　\n\n'),
    ('当事者（ワーカー）', 'タイミーID：　　　　　　（氏名は記録不要）'),
    ('現場責任者',        '　　　　　　　　　　　'),
    ('事故の概要',        '\n\n'),
    ('負傷の程度',        '□軽傷（自力通院）　□重傷（救急搬送）　□物損のみ'),
    ('応急処置の内容',    '　　　　　　　　　　　　　　　　　　　　'),
    ('タイミーへの連絡',  '□連絡済み（時刻：　　　　）　□未連絡'),
    ('再発防止策',        '\n\n'),
    ('報告者署名',        '　　　　　　　　　　　　　　　　'),
]
f2tbl = doc.add_table(rows=len(form2_rows), cols=2)
f2tbl.style = 'Table Grid'
for i, (label, val) in enumerate(form2_rows):
    c0, c1 = f2tbl.rows[i].cells
    c0.width = Cm(4.5)
    c1.width = Cm(11.0)
    set_cell_bg(c0, 'D9E1F2')
    r0 = c0.paragraphs[0].add_run(label)
    r0.bold = True; r0.font.size = Pt(9)
    r0.font.color.rgb = RGBColor(0x1A, 0x37, 0x6E)
    r1 = c1.paragraphs[0].add_run(val)
    r1.font.size = Pt(9)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

out = '/home/user/sacure/サキュレ_タイミー活用社内規程.docx'
doc.save(out)
print('saved:', out)
