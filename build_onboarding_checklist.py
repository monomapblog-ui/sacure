from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ページ設定（A4）
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# デフォルトフォント
style = doc.styles['Normal']
style.font.name = '游ゴシック'
style.font.size = Pt(9)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '游ゴシック')

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_table(doc, headers, rows, col_widths, header_bg='2E74B5'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ヘッダ行
    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # データ行
    timing_bg = {
        '【入社前】':        'D9E1F2',
        '【入社初日】':      'E2EFDA',
        '【入社1週間以内】': 'FFF2CC',
        '【入社1ヶ月以内】': 'FCE4D6',
    }

    for row_data in rows:
        row = table.add_row()
        for i, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[i]
            cell.width = Cm(w)
            p = cell.paragraphs[0]

            # タイミング列（index 1）で背景色
            timing_key = row_data[1] if len(row_data) > 1 else ''
            bg = timing_bg.get(timing_key, 'FFFFFF')
            set_cell_bg(cell, bg)

            # チェック列は中央寄せ
            if i == len(row_data) - 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            if '⚠️' in str(val):
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table

# =====================
# タイトル
# =====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('新入社員 入社受け入れチェックリスト')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('株式会社サキュレ　人事・労務管理用')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
sub.paragraph_format.space_after = Pt(8)

# 凡例
legend = doc.add_paragraph()
legend.alignment = WD_ALIGN_PARAGRAPH.LEFT
lr = legend.add_run(
    '【背景色凡例】  '
    '■ 入社前（青）　■ 入社初日（緑）　■ 1週間以内（黄）　■ 1ヶ月以内（橙）　'
    '　⚠️ = 法的義務・コンプライアンス上、特に重要な項目'
)
lr.font.size = Pt(8)
lr.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
legend.paragraph_format.space_after = Pt(6)

# =====================
# 1. 事務・契約関連
# =====================
add_heading(doc, '1. 事務・契約関連', level=2)

headers1 = ['#', 'タイミング', '項目', '担当', '完了']
col1     = [0.6, 2.6, 8.5, 2.5, 0.8]
rows1 = [
    [1,  '【入社前】',        '雇用契約書の締結・署名捺印',                                    '人事 / 本人', '□'],
    [2,  '【入社前】',        '労働条件通知書の交付・確認',                                    '人事',        '□'],
    [3,  '【入社前】',        '必要書類の提出依頼（マイナンバー、年金手帳、源泉徴収票 等）',  '人事',        '□'],
    [4,  '【入社前】',        '給与振込口座の登録',                                            '本人 / 人事', '□'],
    [5,  '【入社初日】',      '健康保険・厚生年金の加入手続き',                                '人事',        '□'],
    [6,  '【入社初日】',      '雇用保険・労災保険の加入手続き',                                '人事',        '□'],
    [7,  '【入社初日】',      '扶養控除等申告書・住所申告書の記入',                            '本人',        '□'],
    [8,  '【入社1週間以内】', '通勤経路・交通費申請の確認・承認',                              '管理職 / 人事','□'],
    [9,  '【入社1週間以内】', '就業規則・各種規程（服務規律、情報管理規程等）の説明・交付',  '人事',        '□'],
    [10, '【入社1ヶ月以内】', '健康診断の受診案内（必要に応じて）',                            '人事',        '□'],
    [11, '【入社1ヶ月以内】', '人事管理システムへの本人情報登録完了確認',                      '人事',        '□'],
]
add_table(doc, headers1, rows1, col1)

# =====================
# 2. IT・セキュリティ環境
# =====================
add_heading(doc, '2. IT・セキュリティ環境', level=2)

headers2 = ['#', 'タイミング', '項目', '担当', '完了']
col2     = [0.6, 2.6, 8.5, 2.5, 0.8]
rows2 = [
    [12, '【入社前】',        'PC・端末の準備・キッティング（OS設定、MDM登録）',               '情報システム / 管理職','□'],
    [13, '【入社初日】',      '会社メールアドレスの発行・初期パスワード設定',                  '情報システム',         '□'],
    [14, '【入社初日】',      'Slackアカウント発行・必須チャンネルへの招待',                   '情報システム / 管理職','□'],
    [15, '【入社初日】',      '社内SaaS「ラクハイ」へのアクセス権限設定・初期説明',            '情報システム / 管理職','□'],
    [16, '【入社初日】',      '各業務システムアカウント発行（経費精算・勤怠管理・Googleドライブ等）','情報システム',  '□'],
    [17, '【入社初日】',      '⚠️ 情報セキュリティ誓約書の締結',                              '人事 / 本人',          '□'],
    [18, '【入社1週間以内】', 'セキュリティ研修の受講（パスワード管理・フィッシング対策・情報漏洩防止）','本人 / 情報システム','□'],
    [19, '【入社1週間以内】', '二要素認証（2FA）の設定確認',                                  '本人 / 情報システム',  '□'],
    [20, '【入社1週間以内】', 'PC紛失・盗難時の報告フロー周知',                               '情報システム',         '□'],
    [21, '【入社1ヶ月以内】', 'セキュリティ研修の修了確認・テスト実施',                        '情報システム / 人事',  '□'],
    [22, '【入社1ヶ月以内】', '退職者アカウント失効フローの理解確認（知識として）',             '情報システム',         '□'],
]
add_table(doc, headers2, rows2, col2)

# =====================
# 3. 現場・業務関連
# =====================
add_heading(doc, '3. 現場・業務関連（産業廃棄物・物流）', level=2)

headers3 = ['#', 'タイミング', '項目', '担当', '完了']
col3     = [0.6, 2.6, 8.5, 2.5, 0.8]
rows3 = [
    [23, '【入社前】',        '⚠️ 廃棄物処理法の基礎知識資料の事前送付・読み込み',            '人事 / 本人',          '□'],
    [24, '【入社初日】',      '⚠️ 安全教育の実施（KY活動・ヒヤリハット報告制度の説明）',      '管理職',               '□'],
    [25, '【入社初日】',      '⚠️ 産業廃棄物収集運搬許可証の概要説明・携行義務の周知',        '管理職',               '□'],
    [26, '【入社初日】',      '車両使用ルール・日常点検手順の説明（運転業務者対象）',          '管理職',               '□'],
    [27, '【入社初日】',      '⚠️ マニフェスト（管理票）の記載・運用ルール説明',              '管理職 / 先輩社員',     '□'],
    [28, '【入社1週間以内】', '取り扱いマニュアル（廃棄物種別・積載ルール）の読み込み完了確認','管理職 / 本人',        '□'],
    [29, '【入社1週間以内】', '⚠️ 緊急時対応フロー（事故・不法投棄・漏洩）の周知',           '管理職',               '□'],
    [30, '【入社1週間以内】', 'ルート・担当エリアの引き継ぎ・同乗研修',                        '先輩社員 / 管理職',    '□'],
    [31, '【入社1週間以内】', '各拠点（本社・埼玉・板橋・新横浜）の連絡体制確認',              '管理職',               '□'],
    [32, '【入社1ヶ月以内】', '⚠️ コンプライアンス研修受講（下請法、環境法令、個人情報保護）','本人 / 人事',          '□'],
    [33, '【入社1ヶ月以内】', '独り立ち前の業務確認テスト・OJT評価',                          '管理職',               '□'],
]
add_table(doc, headers3, rows3, col3)

# =====================
# 4. オンボーディング
# =====================
add_heading(doc, '4. オンボーディング', level=2)

headers4 = ['#', 'タイミング', '項目', '担当', '完了']
col4     = [0.6, 2.6, 8.5, 2.5, 0.8]
rows4 = [
    [34, '【入社前】',        '入社案内メール送付（初日の集合場所・持ち物・スケジュール）',    '人事',              '□'],
    [35, '【入社初日】',      '社内見学・各部門メンバーへの紹介',                              '管理職',            '□'],
    [36, '【入社初日】',      '会社理念・ビジョン・事業内容説明（ラクハイ含む）',              '管理職 / 社長',     '□'],
    [37, '【入社初日】',      '組織図・役割分担の説明',                                        '管理職',            '□'],
    [38, '【入社初日】',      '直属上司・メンターの確認',                                      '管理職',            '□'],
    [39, '【入社1週間以内】', '業務目標・期待役割の設定面談',                                  '管理職 / 本人',     '□'],
    [40, '【入社1週間以内】', '1on1ミーティングの初回実施',                                    '管理職',            '□'],
    [41, '【入社1週間以内】', '社内コミュニケーションルール（Slack運用・報連相フロー）の説明', '管理職',            '□'],
    [42, '【入社1ヶ月以内】', '30日フォローアップ面談（困りごと・疑問の洗い出し）',            '管理職 / 人事',     '□'],
    [43, '【入社1ヶ月以内】', 'チェックリスト全体の完了確認・人事ファイルへの保管',            '人事',              '□'],
]
add_table(doc, headers4, rows4, col4)

# =====================
# 担当区分説明
# =====================
doc.add_paragraph()
add_heading(doc, '担当区分と運用メモ', level=2)

note_table = doc.add_table(rows=6, cols=2)
note_table.style = 'Table Grid'
note_data = [
    ('人事', '人事・労務担当者が主導'),
    ('管理職', '直属の上司・部門長が主導'),
    ('情報システム', 'IT管理者が主導'),
    ('先輩社員', 'OJT担当の先輩が主導'),
    ('本人', '入社する本人が対応'),
    ('⚠️', '法的義務・コンプライアンス上、特に重要な項目（未完了のまま現場業務開始不可）'),
]
for i, (k, v) in enumerate(note_data):
    c0, c1 = note_table.rows[i].cells
    set_cell_bg(c0, 'D9E1F2')
    c0.width = Cm(3.0)
    c1.width = Cm(11.5)
    r0 = c0.paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.size = Pt(8.5)
    r1 = c1.paragraphs[0].add_run(v)
    r1.font.size = Pt(8.5)
    if k == '⚠️':
        r0.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        r1.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.add_paragraph()
notice = doc.add_paragraph()
nr = notice.add_run(
    '【運用上の注意】　⚠️マークの項目（マニフェスト運用・許可証携行・安全教育・セキュリティ誓約）は、'
    '未完了のまま現場業務を開始させないこと。完了日を必ず記録し、人事ファイルに原本を保管してください。'
)
nr.font.size = Pt(8.5)
nr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
nr.bold = True

out = '/home/user/sacure/サキュレ_入社受け入れチェックリスト.docx'
doc.save(out)
print('saved:', out)
